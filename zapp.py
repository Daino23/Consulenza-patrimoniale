import streamlit as st
from docx import Document
from datetime import date
from io import BytesIO

st.set_page_config(page_title="Consulenza Patrimoniale - Studio Dainotti", layout="wide")
st.title("Scheda Consulenza Patrimoniale - Studio Dainotti")

def init_state():
    if "tab_index" not in st.session_state:
        st.session_state.tab_index = 0
    if "responses" not in st.session_state:
        st.session_state.responses = {}
    if "familiari_count" not in st.session_state:
        st.session_state.familiari_count = 0
    if "patrimonio_count" not in st.session_state:
        st.session_state.patrimonio_count = 0
    if "debiti_count" not in st.session_state:
        st.session_state.debiti_count = 0
    if "obiettivi_count" not in st.session_state:
        st.session_state.obiettivi_count = 0

init_state()

custom_button_style = """
    <style>
    .prosegui-button button {
        position: relative;
        float: right;
        margin-top: 2em;
        font-weight: bold;
        padding: 0.6em 1.2em;
        border-radius: 6px;
        border: 1px solid #ccc;
        background-color: #f0f0f0;
        color: black;
    }
    </style>
"""

st.markdown(custom_button_style, unsafe_allow_html=True)


tabs = ["🏠 Famiglia", "💼 Patrimonio", "💳 Debiti", "🎯 Obiettivi", "📄 Documento"]
current_tab = st.session_state.tab_index
tab_objects = st.tabs(tabs)

with tab_objects[0]:
    st.header("🏠 Famiglia e situazione personale")
    st.session_state.responses["Nome e cognome"] = st.text_input("Nome e cognome")
    st.session_state.responses["Data e luogo di nascita"] = st.text_input("Data e luogo di nascita")
    st.session_state.responses["Codice fiscale"] = st.text_input("Codice fiscale")
    st.session_state.responses["Indirizzo"] = st.text_input("Indirizzo")
    stato_civile = st.selectbox("Stato civile", ["", "Celibe/Nubile", "Coniugato/a", "Separato/a", "Divorziato/a", "Vedovo/a"])
    st.session_state.responses["Stato civile"] = stato_civile

    if stato_civile == "Coniugato/a":
        st.session_state.responses["Coniuge"] = st.text_input("Nome e cognome del coniuge")

    if st.checkbox("Hai figli?"):
        num_figli = st.number_input("Numero di figli", min_value=1, step=1)
        figli = []
        for i in range(int(num_figli)):
            with st.expander(f"Figlio #{i+1}"):
                nome = st.text_input(f"Nome figlio #{i+1}", key=f"figlio_nome_{i}")
                nascita = st.text_input(f"Data di nascita figlio #{i+1}", key=f"figlio_nascita_{i}")
                codice = st.text_input(f"Codice fiscale figlio #{i+1}", key=f"figlio_cf_{i}")
                figli.append(f"{nome} - {nascita} - {codice}")
        st.session_state.responses["Figli"] = figli

    st.subheader("👥 Altri familiari a carico")
    if st.button("➕ Aggiungi familiare a carico"):
        st.session_state.familiari_count += 1

    familiari = []
    for i in range(st.session_state.familiari_count):
        with st.expander(f"Familiare #{i+1}"):
            nome = st.text_input(f"Nome #{i+1}", key=f"fam_nome_{i}")
            relazione = st.text_input(f"Relazione #{i+1}", key=f"fam_rel_{i}")
            codice = st.text_input(f"Codice fiscale #{i+1}", key=f"fam_cf_{i}")
            familiari.append(f"{nome} ({relazione}) - CF: {codice}")
    st.session_state.responses["Altri familiari a carico"] = familiari

with tab_objects[1]:
    st.header("💼 Area patrimoniale")
    if st.button("➕ Aggiungi bene patrimoniale"):
        st.session_state.patrimonio_count += 1

    patrimonio = []
    for i in range(st.session_state.patrimonio_count):
        with st.expander(f"Bene #{i+1}"):
            tipo = st.selectbox(f"Tipo di bene #{i+1}", ["Immobile", "Conto corrente", "Fondo", "ETF", "Trust", "Altro"], key=f"tipo_{i}")
            descrizione = st.text_input(f"Descrizione #{i+1}", key=f"desc_{i}")
            intestatario = st.text_input(f"Intestatario #{i+1}", key=f"intest_{i}")
            valore = st.number_input(f"Valore stimato € #{i+1}", min_value=0.0, step=100.0, key=f"valore_{i}")
            patrimonio.append(f"{tipo} - {descrizione} - {intestatario} - {valore:.2f} €")
    st.session_state.responses["Situazione patrimoniale"] = patrimonio

with tab_objects[2]:
    st.header("💳 Debiti ricorrenti")
    if st.button("➕ Aggiungi debito"):
        st.session_state.debiti_count += 1

    debiti = []
    for i in range(st.session_state.debiti_count):
        with st.expander(f"Debito #{i+1}"):
            tipo = st.text_input(f"Tipo #{i+1}", key=f"deb_tipo_{i}")
            importo = st.number_input(f"Importo mensile € #{i+1}", min_value=0.0, step=10.0, key=f"deb_importo_{i}")
            debiti.append(f"{tipo} - {importo:.2f} € / mese")
    st.session_state.responses["Debiti ricorrenti"] = debiti

with tab_objects[3]:
    st.header("🎯 Obiettivi economici")

    if st.button("➕ Aggiungi obiettivo"):
        st.session_state.obiettivi_count += 1

    obiettivi = []
    for i in range(st.session_state.obiettivi_count):
        with st.expander(f"Obiettivo #{i+1}"):
            descrizione = st.text_input(f"Descrizione #{i+1}", key=f"ob_desc_{i}")
            importo = st.number_input(f"Importo da raggiungere (€) #{i+1}", min_value=0.0, step=500.0, key=f"ob_importo_{i}")
            tempo = st.text_input(f"Tempo previsto (es. 12 mesi) #{i+1}", key=f"ob_tempo_{i}")
            obiettivi.append(f"{descrizione} | {importo:.2f} € entro {tempo}")
    st.session_state.responses["Obiettivi economici"] = obiettivi

with tab_objects[4]:
    st.header("📄 Genera documento Word")

    if st.button("Crea documento Word"):
        doc = Document()
        doc.add_heading("Scheda Consulenza Patrimoniale", 0)
        doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}")

        for sezione, risposte in st.session_state.responses.items():
            doc.add_heading(sezione, level=1)
            if isinstance(risposte, list):
                for voce in risposte:
                    doc.add_paragraph(str(voce), style='List Bullet')
            else:
                doc.add_paragraph(str(risposte))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Scarica documento Word",
            data=buffer,
            file_name="scheda_consulenza_patrimoniale.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if st.button("📊 Scarica file Excel"):
        import pandas as pd

        patrimonio_data = []
        for i in range(st.session_state.patrimonio_count):
            tipo = st.session_state.get(f"tipo_{i}", "")
            descrizione = st.session_state.get(f"desc_{i}", "")
            intestatario = st.session_state.get(f"intest_{i}", "")
            valore = st.session_state.get(f"valore_{i}", 0.0)
            patrimonio_data.append({
                "Tipo": tipo,
                "Descrizione": descrizione,
                "Intestatario": intestatario,
                "Valore (€)": valore
            })

        debiti_data = []
        for i in range(st.session_state.debiti_count):
            tipo = st.session_state.get(f"deb_tipo_{i}", "")
            importo = st.session_state.get(f"deb_importo_{i}", 0.0)
            debiti_data.append({
                "Tipo": tipo,
                "Importo mensile (€)": importo
            })

        obiettivi_data = []
        for i in range(st.session_state.obiettivi_count):
            descrizione = st.session_state.get(f"ob_desc_{i}", "")
            importo = st.session_state.get(f"ob_importo_{i}", 0.0)
            tempo = st.session_state.get(f"ob_tempo_{i}", "")
            obiettivi_data.append({
                "Obiettivo": descrizione,
                "Importo (€)": importo,
                "Tempo previsto": tempo
            })

        with BytesIO() as output:
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                pd.DataFrame(patrimonio_data).to_excel(writer, index=False, sheet_name="Patrimonio")
                pd.DataFrame(debiti_data).to_excel(writer, index=False, sheet_name="Debiti")
                pd.DataFrame(obiettivi_data).to_excel(writer, index=False, sheet_name="Obiettivi")
            output.seek(0)

            st.download_button(
                label="📊 Scarica file Excel",
                data=output,
                file_name="dati_consulenza.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
