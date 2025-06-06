import streamlit as st
from docx import Document
from datetime import date
from io import BytesIO

st.set_page_config(page_title="Consulenza Patrimoniale - Studio Dainotti", layout="wide")
st.title("Scheda Consulenza Patrimoniale - Studio Dainotti")

def init_state():
    if "responses" not in st.session_state:
        st.session_state.responses = {}

init_state()

sections = {
    "ğŸ  Famiglia e situazione personale": [],
    "ğŸ’¼ Area patrimoniale": [],
    "âš–ï¸ Pianificazione successoria e protezione": [
        "Nome erede, data nascita, parentela, situazione patrimoniale, note",
        "Donazioni previste", "Beni da destinare specificamente", "Diritti da riservare", "Testamento previsto", "Obiettivo evitare conflitti",
        "Testamento esistente", "Donazioni giÃ  effettuate", "Clausole o intestazioni", "Trust o fondi patrimoniali", "Quote aziendali"
    ],
    "ğŸ” Obiettivi e bisogni": [
        "Protezione del patrimonio", "Ottimizzazione fiscale", "Passaggio generazionale", "Altro",
        "Contenziosi", "FragilitÃ ", "Beni allâ€™estero", "Altre criticitÃ "
    ],
    "ğŸ“ Allegati richiesti": [
        "Visure catastali", "Atti notarili", "Documentazione assicurativa", "Estratti conto", "Testamenti o donazioni", "Quote societarie"
    ]
}

tab_titles = list(sections.keys()) + ["ğŸ“„ Esporta documento"]
tabs = st.tabs(tab_titles)

# TAB 0 - Famiglia
with tabs[0]:
    st.header("ğŸ  Famiglia e situazione personale")
    st.session_state.responses["Nome e cognome"] = st.text_input("Nome e cognome")
    st.session_state.responses["Data e luogo di nascita"] = st.text_input("Data e luogo di nascita")
    st.session_state.responses["Codice fiscale"] = st.text_input("Codice fiscale")
    st.session_state.responses["Indirizzo"] = st.text_input("Indirizzo")
    stato_civile = st.selectbox("Stato civile", ["", "Celibe/Nubile", "Coniugato/a", "Separato/a", "Divorziato/a", "Vedovo/a"])
    st.session_state.responses["Stato civile"] = stato_civile

    if stato_civile == "Coniugato/a":
        st.subheader("ğŸ‘« Dati del coniuge")
        st.session_state.responses["Coniuge"] = st.text_input("Nome e cognome del coniuge")

    ha_figli = st.checkbox("Hai figli?")
    figli = []
    if ha_figli:
        num_figli = st.number_input("Quanti figli vuoi inserire?", min_value=1, step=1)
        for i in range(int(num_figli)):
            with st.expander(f"Figlio #{i+1}"):
                nome = st.text_input(f"Nome figlio #{i+1}", key=f"figlio_nome_{i}")
                nascita = st.text_input(f"Data di nascita figlio #{i+1}", key=f"figlio_nascita_{i}")
                codice = st.text_input(f"Codice fiscale figlio #{i+1}", key=f"figlio_cf_{i}")
                figli.append(f"{nome} - {nascita} - {codice}")
    st.session_state.responses["Figli"] = figli

    st.subheader("ğŸ‘¥ Altri familiari a carico")
    if "familiari_count" not in st.session_state:
        st.session_state.familiari_count = 0
        st.session_state.familiari = []

    if st.button("â• Aggiungi familiare a carico"):
        st.session_state.familiari_count += 1

    familiari_input = []
    for i in range(st.session_state.familiari_count):
        with st.expander(f"Familiare a carico #{i+1}"):
            parentela = st.selectbox(f"Parentela #{i+1}", ["Genitore", "Fratello/Sorella", "Altro"], key=f"parentela_{i}")
            nome = st.text_input(f"Nome familiare #{i+1}", key=f"fam_nome_{i}")
            nascita = st.text_input(f"Data di nascita #{i+1}", key=f"fam_nascita_{i}")
            codice = st.text_input(f"Codice fiscale #{i+1}", key=f"fam_cf_{i}")
            note = st.text_area(f"Note #{i+1}", key=f"fam_note_{i}")
            familiari_input.append(f"{parentela}: {nome} - {nascita} - {codice}. Note: {note}")
    st.session_state.responses["Altri familiari a carico"] = familiari_input

    st.subheader("ğŸ’¼ Situazione lavorativa e reddituale")
    st.session_state.responses["Occupazione"] = st.text_input("Occupazione")
    st.session_state.responses["Reddito lordo annuo"] = st.text_input("Reddito lordo annuo")
    st.session_state.responses["Altri redditi"] = st.text_area("Altri redditi")

    st.subheader("ğŸ’³ Debiti ricorrenti")

    if "debiti_count" not in st.session_state:
        st.session_state.debiti_count = 0
        st.session_state.debiti = []

    if st.button("â• Aggiungi debito ricorrente"):
        st.session_state.debiti_count += 1

    debiti_input = []
    for i in range(st.session_state.debiti_count):
        with st.expander(f"Debito ricorrente #{i+1}"):
            tipo = st.text_input(f"Tipo di debito #{i+1}", key=f"tipo_debito_{i}")
            importo = st.number_input(f"Importo mensile in euro #{i+1}", min_value=0.0, step=10.0, key=f"importo_debito_{i}")
            debiti_input.append(f"{tipo} - {importo:.2f} â‚¬ / mese")
    st.session_state.responses["Debiti ricorrenti"] = debiti_input

# TAB 1 - Patrimonio
with tabs[1]:
    st.header("ğŸ’¼ Area patrimoniale")
    if "patrimonio_count" not in st.session_state:
        st.session_state.patrimonio_count = 0
        st.session_state.patrimonio_voci = []

    if st.button("â• Aggiungi voce patrimoniale"):
        st.session_state.patrimonio_count += 1

    patrimonio = []
    for i in range(st.session_state.patrimonio_count):
        with st.expander(f"Voce patrimoniale #{i+1}"):
            tipo = st.selectbox(f"Tipo di patrimonio #{i+1}", [
                "Immobile", "Conto corrente", "Fondo comune", "ETF", "Trust", "Polizza vita", "Quota aziendale", "Altro"
            ], key=f"tipo_patr_{i}")
            descrizione = st.text_input(f"Descrizione sintetica #{i+1}", key=f"desc_patr_{i}")
            intestatario = st.text_input(f"Intestatario #{i+1}", key=f"intestatario_patr_{i}")
            note = st.text_area(f"Note o dettagli aggiuntivi #{i+1}", key=f"note_patr_{i}")

            if tipo in ["Conto corrente", "Fondo comune", "ETF", "Trust", "Polizza vita", "Quota aziendale"]:
                valore = st.number_input(f"Valore stimato in euro #{i+1}", min_value=0.0, step=1000.0, key=f"valore_patr_{i}")
                riga = f"{tipo} - {descrizione} | Valore: {valore:.2f} â‚¬ | Intestatario: {intestatario} | Note: {note}"
            else:
                riga = f"{tipo} - {descrizione} | Intestatario: {intestatario} | Note: {note}"
            patrimonio.append(riga)

    st.session_state.responses["Patrimonio dettagliato"] = patrimonio

# TAB 2 - Successoria
with tabs[2]:
    st.header("âš–ï¸ Pianificazione successoria e protezione")
    for question in sections["âš–ï¸ Pianificazione successoria e protezione"]:
        response = st.text_area(question, key=question)
        st.session_state.responses[question] = response

# TAB 3 - Obiettivi
with tabs[3]:
    st.header("ğŸ” Obiettivi e bisogni")
    for question in sections["ğŸ” Obiettivi e bisogni"]:
        response = st.text_area(question, key=question)
        st.session_state.responses[question] = response

# TAB 4 - Allegati
with tabs[4]:
    st.header("ğŸ“ Allegati richiesti")
    for question in sections["ğŸ“ Allegati richiesti"]:
        response = st.text_area(question, key=question)
        st.session_state.responses[question] = response

# TAB 5 - Esporta
with tabs[5]:
    st.header("ğŸ“„ Esporta il documento Word finale")

    st.subheader("ğŸ” Riepilogo delle risposte")
    for key, value in st.session_state.responses.items():
        if isinstance(value, list):
            if value:
                st.markdown(f"**{key}:**")
                for v in value:
                    st.markdown(f"- {v}")
        elif value:
            st.markdown(f"**{key}:** {value}")

    if st.button("Genera documento Word"):
        doc = Document()
        doc.add_paragraph("Studio Dainotti - Consulenza patrimoniale, tributaria e del credito\nVia Roma 52, Porto Valtravaglia (VA) - www.dainotti.com - info@dainotti.com")
        doc.add_paragraph(f"Data compilazione: {date.today().strftime('%d/%m/%Y')}")
        doc.add_heading("Scheda Raccolta Informazioni - Consulenza Patrimoniale", 0)

        for section_title in sections:
            doc.add_heading(section_title, level=1)
            for key in st.session_state.responses:
                if key in sections[section_title] or section_title == "ğŸ  Famiglia e situazione personale":
                    value = st.session_state.responses[key]
                    if isinstance(value, list):
                        for item in value:
                            doc.add_paragraph(f"- {item}")
                    else:
                        doc.add_paragraph(f"{key}: {value}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("Documento generato con successo!")
        st.download_button(
            label="ğŸ“¥ Scarica il file Word",
            data=buffer,
            file_name="Scheda_Consulenza_Patrimoniale.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
