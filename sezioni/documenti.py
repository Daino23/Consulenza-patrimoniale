import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import datetime
from docx import Document # Importa la libreria python-docx
from docx.shared import Inches # Per immagini, se volessimo aggiungerne

def sezione_documenti():
    st.subheader("📄 Generazione Documenti")
    st.write("Genera un file Excel o Word con i dati inseriti nelle altre sezioni.")

    # --- Recupero dei dati da st.session_state per entrambe le esportazioni ---
    cliente_data = st.session_state.get('cliente_data', {})
    
    # Verifica se ci sono dati per l'esportazione
    has_data_to_export = False
    if cliente_data.get('dati_personali') or \
       (cliente_data.get('dati_familiari') and (cliente_data['dati_familiari'].get('coniuge') or cliente_data['dati_familiari']['figli'] or cliente_data['dati_familiari']['altri_familiari'])) or \
       cliente_data.get('profilo_finanziario'): # Aggiunto controllo per profilo finanziario
        has_data_to_export = True
    if st.session_state.get('beni_patrimoniali'):
        has_data_to_export = True
    if st.session_state.get('lista_debiti'):
        has_data_to_export = True
    if st.session_state.get('lista_obiettivi'):
        has_data_to_export = True
    # Aggiungi qui altre condizioni se hai altre sezioni con dati esportabili (es. CRM)

    # --- Sezione di Esportazione Excel ---
    if has_data_to_export:
        output_excel = BytesIO() # Usiamo un BytesIO separato per Excel
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            # Dati Cliente Anagrafica
            if cliente_data.get('dati_personali'):
                df_cliente_personale = pd.DataFrame([cliente_data['dati_personali']])
                df_cliente_personale.to_excel(writer, index=False, sheet_name="Cliente_DatiPersonali")

            if cliente_data.get('profilo_finanziario'):
                df_cliente_profilo = pd.DataFrame([cliente_data['profilo_finanziario']])
                df_cliente_profilo.to_excel(writer, index=False, sheet_name="Cliente_ProfiloFin")

            if cliente_data.get('dati_familiari'):
                familiari_list = []
                if cliente_data['dati_familiari'].get('coniuge'):
                    familiari_list.append({"Tipo": "Coniuge", **cliente_data['dati_familiari']['coniuge']})
                if cliente_data['dati_familiari'].get('figli'):
                    for figlio in cliente_data['dati_familiari']['figli']:
                        familiari_list.append({"Tipo": "Figlio", **figlio})
                if cliente_data['dati_familiari'].get('altri_familiari'):
                    for alt_fam in cliente_data['dati_familiari']['altri_familiari']:
                        familiari_list.append({"Tipo": "Altro Familiare", **alt_fam})

                if familiari_list:
                    df_familiari = pd.DataFrame(familiari_list)
                    df_familiari.to_excel(writer, index=False, sheet_name="Cliente_Familiari")

            # Dati Patrimonio
            if st.session_state.get("beni_patrimoniali"):
                df_patrimonio = pd.DataFrame(st.session_state.beni_patrimoniali)
                df_patrimonio.to_excel(writer, index=False, sheet_name="Patrimonio")

            # Dati Debiti
            if st.session_state.get("lista_debiti"):
                df_debiti = pd.DataFrame(st.session_state.lista_debiti)
                df_debiti.to_excel(writer, index=False, sheet_name="Debiti")

            # Dati Obiettivi
            if st.session_state.get("lista_obiettivi"):
                df_obiettivi = pd.DataFrame(st.session_state.lista_obiettivi)
                df_obiettivi.to_excel(writer, index=False, sheet_name="Obiettivi")

        output_excel.seek(0)
        st.download_button(
            label="📥 Scarica Riepilogo Completo in Excel",
            data=output_excel.getvalue(),
            file_name="riepilogo_consulenza.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_excel_button" # Chiave unica per il bottone
        )
    else:
        st.info("Compila le sezioni Anagrafica Cliente, Patrimonio, Debiti o Obiettivi per generare i documenti.")

    st.markdown("---") # Linea separatrice

    # --- Sezione di Esportazione Word ---
    st.subheader("📝 Generazione Report in Word")
    st.write("Genera un documento Word con un riepilogo testuale delle informazioni del cliente.")

    if has_data_to_export:
        if st.button("📄 Genera Report in Word", key="generate_word_button"): # Bottone per avviare la generazione
            document = Document()
            document.add_heading("Report di Consulenza Patrimoniale", level=1)
            document.add_paragraph(f"Data del Report: {datetime.now().strftime('%d/%m/%Y')}")
            document.add_paragraph("---")

            # Sezione Dati Anagrafici
            document.add_heading("1. Anagrafica Cliente", level=2)
            if cliente_data.get('dati_personali'):
                dp = cliente_data['dati_personali']
                document.add_paragraph(f"Nome e Cognome: {dp.get('Nome e cognome', 'N/A')}")
                document.add_paragraph(f"Data e Luogo di Nascita: {dp.get('Data e luogo di nascita', 'N/A')}")
                document.add_paragraph(f"Codice Fiscale: {dp.get('Codice fiscale', 'N/A')}")
                document.add_paragraph(f"Indirizzo: {dp.get('Indirizzo', 'N/A')}")
                document.add_paragraph(f"Stato Civile: {dp.get('Stato civile', 'N/A')}")
            else:
                document.add_paragraph("Nessun dato personale inserito.")

            if cliente_data.get('dati_familiari'):
                dfam = cliente_data['dati_familiari']
                if dfam.get('coniuge'):
                    document.add_paragraph(f"Coniuge: {dfam['coniuge'].get('Nome e cognome', 'N/A')}")
                if dfam.get('figli'):
                    document.add_paragraph("Figli:")
                    for figlio in dfam['figli']:
                        document.add_paragraph(f"- {filho.get('Nome', 'N/A')} (Nato il: {filho.get('Data di nascita', 'N/A')}, CF: {filho.get('Codice fiscale', 'N/A')})")
                if dfam.get('altri_familiari'):
                    document.add_paragraph("Altri Familiari a Carico:")
                    for alt_fam in dfam['altri_familiari']:
                        document.add_paragraph(f"- {alt_fam.get('Nome', 'N/A')} (Relazione: {alt_fam.get('Relazione', 'N/A')}, CF: {alt_fam.get('Codice fiscale', 'N/A')})")
            
            if cliente_data.get('profilo_finanziario'):
                pf = cliente_data['profilo_finanziario']
                document.add_heading("Profilo Finanziario", level=3)
                document.add_paragraph(f"Professione: {pf.get('Professione', 'N/A')}")
                document.add_paragraph(f"Reddito Annuo Netto: {pf.get('Reddito annuo netto', 'N/A')} €")
                document.add_paragraph(f"Patrimonio Liquido: {pf.get('Patrimonio liquido', 'N/A')} €")
                document.add_paragraph(f"Spese Mensili: {pf.get('Spese mensili', 'N/A')} €")
                document.add_paragraph(f"Tolleranza al Rischio: {pf.get('Tolleranza al rischio', 'N/A')}")

            # Sezione Patrimonio
            document.add_heading("2. Patrimonio", level=2)
            if st.session_state.get("beni_patrimoniali"):
                for i, bene in enumerate(st.session_state.beni_patrimoniali):
                    document.add_paragraph(f"- {bene.get('Tipo', 'N/A')}: {bene.get('Descrizione', 'N/A')} (Intestatario: {bene.get('Intestatario', 'N/A')}, Valore: {bene.get('Valore', 0.0):.2f} €)")
            else:
                document.add_paragraph("Nessun bene patrimoniale inserito.")

            # Sezione Debiti
            document.add_heading("3. Debiti", level=2)
            if st.session_state.get("lista_debiti"):
                for i, debito in enumerate(st.session_state.lista_debiti):
                    document.add_paragraph(f"- {debito.get('Tipo', 'N/A')}: Importo Mensile {debito.get('Importo Mensile', 0.0):.2f} €")
            else:
                document.add_paragraph("Nessun debito inserito.")

            # Sezione Obiettivi
            document.add_heading("4. Obiettivi Economici", level=2)
            if st.session_state.get("lista_obiettivi"):
                for i, obiettivo in enumerate(st.session_state.lista_obiettivi):
                    document.add_paragraph(f"- {obiettivo.get('Descrizione', 'N/A')}: Target {obiettivo.get('Importo Target', 0.0):.2f} € (Tempo previsto: {obiettivo.get('Tempo Previsto', 'N/A')})")
            else:
                document.add_paragraph("Nessun obiettivo economico inserito.")

            # Salva il documento in un BytesIO
            output_word = BytesIO()
            document.save(output_word)
            output_word.seek(0)

            st.download_button(
                label="Clicca qui per scaricare il Report Word",
                data=output_word.getvalue(),
                file_name="report_consulenza.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word_button" # Chiave unica per il bottone
            )
            st.success("Report Word generato e pronto per il download!")

    st.markdown("---") # Linea separatrice

    # --- Sezione Calcolatrice ---
    st.subheader("🧮 Calcolatrice Rapida")
    st.markdown("Esegui calcoli semplici direttamente qui.")

    col_calc1, col_calc2 = st.columns(2)

    with col_calc1:
        num1 = st.number_input("Primo Numero", value=0.0, format="%.2f", key="calc_num1")
    with col_calc2:
        num2 = st.number_input("Secondo Numero", value=0.0, format="%.2f", key="calc_num2")

    operation = st.selectbox("Seleziona Operazione", ["+", "-", "*", "/"], key="calc_op")

    result = 0.0
    if st.button("Calcola", key="calc_button"):
        if operation == "+":
            result = num1 + num2
        elif operation == "-":
            result = num1 - num2
        elif operation == "*":
            result = num1 * num2
        elif operation == "/":
            if num2 != 0:
                result = num1 / num2
            else:
                st.error("Errore: Divisione per zero!")
                result = "Non definito" 
        
        if isinstance(result, float):
            st.success(f"Risultato: {result:.2f}")
        else:
            st.success(f"Risultato: {result}")

        st.session_state["calc_result"] = result

    if "calc_result" in st.session_state and st.session_state["calc_result"] != "Non definito":
        st.write(f"Ultimo Risultato Calcolato: **{st.session_state['calc_result']:.2f}**")
    elif "calc_result" in st.session_state and st.session_state["calc_result"] == "Non definito":
         st.write(f"Ultimo Risultato Calcolato: **{st.session_state['calc_result']}**")
